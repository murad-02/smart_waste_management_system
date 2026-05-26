"""Generate a Word document containing UML / architectural diagrams for SWMS.

The diagrams are emitted as Mermaid code blocks (render at https://mermaid.live)
together with explanatory tables and ASCII fallbacks so the document is useful
both for viewing and for regenerating PNG/SVG renders later.
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0E, 0x4D, 0x92)
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_code_block(doc, code_text, language_label=None):
    """Add a monospaced, shaded code block (Mermaid / ASCII)."""
    if language_label:
        cap = doc.add_paragraph()
        run = cap.add_run(f"[{language_label}]")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p


def add_table(doc, header, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------
def build_document(out_path):
    doc = Document()

    # ---- Cover ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Smart Waste Management System (SWMS)")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x00, 0xB8, 0x94)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("System Design Diagrams")
    run.bold = True
    run.font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Class · Use Case · Swim-lane · Sequence · ER · Data Flow Diagrams"
    )
    run.italic = True
    run.font.size = Pt(12)

    doc.add_paragraph()
    add_paragraph(
        doc,
        "This document describes the architecture of the SWMS desktop application "
        "(PyQt5 + SQLAlchemy + SQLite + YOLOv8). It covers the waste-detection "
        "subsystem, role-based authentication, alerts, reports and the fleet & truck "
        "management subsystem. Each diagram is provided as Mermaid source "
        "(render with https://mermaid.live or any Mermaid-aware viewer) and "
        "complemented by a short ASCII / tabular description.",
    )
    doc.add_page_break()

    # ============================================================
    # 1. CLASS DIAGRAM
    # ============================================================
    add_heading(doc, "1. Class Diagram", level=1)
    add_paragraph(
        doc,
        "The class diagram combines (a) the SQLAlchemy ORM models in database/ "
        "and (b) the service classes in core/ together with their key "
        "PyQt5 UI screens. Relationships shown reflect the foreign keys and "
        "back_populates declarations in database/models.py and "
        "database/fleet_models.py.",
    )

    mermaid_class = """classDiagram
    %% ---- Domain Models ----
    class User {
      +int id
      +str username
      +str password_hash
      +str full_name
      +str email
      +str role
      +bool is_active
      +datetime created_at
      +datetime last_login
    }
    class Detection {
      +int id
      +str image_path
      +str result_image_path
      +str waste_category
      +float confidence
      +str bin_fill_level
      +int detected_by
      +str status
      +int verified_by
      +datetime detected_at
      +str notes
    }
    class AlertRule {
      +int id
      +str rule_name
      +str category
      +int threshold_value
      +str period
      +str notify_email
      +bool is_active
      +int created_by
    }
    class Alert {
      +int id
      +int rule_id
      +str message
      +str severity
      +bool is_acknowledged
      +int acknowledged_by
      +datetime triggered_at
    }
    class ActivityLog {
      +int id
      +int user_id
      +str action
      +str details
      +datetime timestamp
    }
    class Report {
      +int id
      +str report_type
      +str file_path
      +int generated_by
      +date date_range_start
      +date date_range_end
    }
    class AppSetting {
      +int id
      +str key
      +str value
    }
    class Truck {
      +int id
      +str truck_code
      +str plate_number
      +float capacity
      +str fuel_type
      +str status
      +str assigned_zone
      +date purchase_date
      +bool is_active
    }
    class Driver {
      +int id
      +str name
      +str license_number
      +int assigned_truck_id
      +str status
    }
    class Route {
      +int id
      +str route_name
      +str zone
      +float estimated_distance
      +int estimated_duration
    }
    class CollectionTrip {
      +int id
      +int truck_id
      +int driver_id
      +int route_id
      +datetime start_time
      +datetime end_time
      +float waste_weight
      +str trip_status
      +int created_by
    }
    class MaintenanceRecord {
      +int id
      +int truck_id
      +str service_type
      +date service_date
      +date next_service_date
      +float cost
    }

    %% ---- Service Layer ----
    class AuthManager {
      +login(username, password)
      +create_user(...)
      +update_user(id, **kwargs)
      +deactivate_user(id, admin_id)
      +check_permission(user, role)
    }
    class DetectionEngine {
      +detect(image_path, user_id)
      +detect_video_stream(path)
      +update_detection_status(id, status, verified_by)
      +export_detections_csv(list, path)
      +export_detections_excel(list, path)
    }
    class AlertManager {
      +create_rule(...)
      +check_alerts() list
      +acknowledge_alert(id, user)
      +send_test_alert(rule_id)
    }
    class NotificationService {
      +send_email(to, subject, body)
      +send_email_verbose(...)
    }
    class AnalyticsEngine {
      +get_today_stats()
      +get_category_distribution()
      +get_daily_counts(days)
      +get_fill_level_distribution()
    }
    class ReportEngine {
      +generate_report(type, start, end, user_id)
      +get_all_reports()
      +delete_report(id)
    }
    class LogManager {
      +log_activity(user_id, action, details)
      +get_logs(user_id, start, end)
    }
    class TruckService {
      +list_trucks()
      +create(actor, data)
      +update(actor, id, data)
      +soft_delete(actor, id)
    }
    class DriverService
    class RouteService
    class TripService
    class MaintenanceService

    %% ---- UI Layer ----
    class MainWindow
    class LoginScreen
    class DashboardScreen
    class DetectionScreen
    class HistoryScreen
    class AlertsScreen
    class ReportsScreen
    class UsersScreen
    class SettingsScreen
    class FleetDashboardScreen
    class TrucksScreen
    class DriversScreen
    class RoutesScreen
    class TripsScreen
    class MaintenanceScreen

    %% ---- Relationships (ORM) ----
    User "1" --o "*" Detection : detects
    User "1" --o "*" Detection : verifies
    User "1" --o "*" ActivityLog : performs
    User "1" --o "*" AlertRule : creates
    User "1" --o "*" Alert : acknowledges
    User "1" --o "*" Report : generates
    AlertRule "1" --o "*" Alert : triggers
    Truck "1" --o "*" Driver : assigned
    Truck "1" --o "*" CollectionTrip : runs
    Truck "1" --o "*" MaintenanceRecord : serviced
    Driver "1" --o "*" CollectionTrip : drives
    Route "1" --o "*" CollectionTrip : follows

    %% ---- Service Dependencies ----
    AlertManager --> NotificationService
    AlertManager --> AlertRule
    AlertManager --> Alert
    AlertManager --> Detection
    DetectionEngine --> Detection
    AuthManager --> User
    ReportEngine --> AnalyticsEngine
    ReportEngine --> Detection
    ReportEngine --> Report
    AnalyticsEngine --> Detection
    AnalyticsEngine --> Alert
    LogManager --> ActivityLog
    TruckService --> Truck
    DriverService --> Driver
    RouteService --> Route
    TripService --> CollectionTrip
    MaintenanceService --> MaintenanceRecord

    %% ---- UI -> Service ----
    MainWindow --> LoginScreen
    MainWindow --> DashboardScreen
    LoginScreen --> AuthManager
    DetectionScreen --> DetectionEngine
    DetectionScreen --> AlertManager
    HistoryScreen --> DetectionEngine
    AlertsScreen --> AlertManager
    ReportsScreen --> ReportEngine
    UsersScreen --> AuthManager
    SettingsScreen --> NotificationService
    DashboardScreen --> AnalyticsEngine
    TrucksScreen --> TruckService
    DriversScreen --> DriverService
    RoutesScreen --> RouteService
    TripsScreen --> TripService
    MaintenanceScreen --> MaintenanceService
"""
    add_code_block(doc, mermaid_class, "Mermaid - classDiagram")

    add_paragraph(doc, "Layer Summary", bold=True, size=12)
    add_table(
        doc,
        ["Layer", "Modules / Classes", "Responsibility"],
        [
            ["UI (PyQt5)",
             "MainWindow, LoginScreen, DashboardScreen, DetectionScreen, "
             "HistoryScreen, AlertsScreen, ReportsScreen, UsersScreen, "
             "SettingsScreen, Fleet*Screen",
             "Render widgets, capture user events, delegate to services"],
            ["Service / Core",
             "AuthManager, DetectionEngine, AlertManager, AnalyticsEngine, "
             "ReportEngine, NotificationService, LogManager, TruckService, "
             "DriverService, RouteService, TripService, MaintenanceService",
             "Business rules, session lifecycle, validation, persistence"],
            ["Persistence (SQLAlchemy)",
             "Base, User, Detection, AlertRule, Alert, ActivityLog, Report, "
             "AppSetting, Truck, Driver, Route, CollectionTrip, "
             "MaintenanceRecord",
             "ORM mapping to SQLite tables"],
            ["External",
             "YOLOv8 (Ultralytics), bcrypt, ReportLab, matplotlib, smtplib, "
             "openpyxl, Pillow / OpenCV",
             "Detection inference, hashing, PDF/Excel export, email, charts"],
        ],
        col_widths=[1.3, 3.0, 3.0],
    )
    doc.add_page_break()

    # ============================================================
    # 2. USE CASE DIAGRAM
    # ============================================================
    add_heading(doc, "2. Use Case Diagram", level=1)
    add_paragraph(
        doc,
        "Actors are derived from the three roles defined in AuthManager "
        "(admin, supervisor, operator) plus the external SMTP server. "
        "Permissions correspond to the ROLE_HIERARCHY in core/auth_manager.py "
        "and the ACTIONS table in core/fleet/fleet_permissions.py.",
    )

    mermaid_usecase = """flowchart LR
    %% Actors
    Operator((Operator))
    Supervisor((Supervisor))
    Admin((Administrator))
    SMTP[[SMTP Server]]

    %% System boundary
    subgraph SWMS [Smart Waste Management System]
      UC1([Login / Logout])
      UC2([Run Waste Detection on Image / Video])
      UC3([View Dashboard Statistics])
      UC4([View Waste History])
      UC5([Verify / Reject Detection])
      UC6([Export Detections CSV/Excel])
      UC7([Acknowledge Alert])
      UC8([Manage Alert Rules])
      UC9([Generate PDF Report])
      UC10([Open / Delete Report])
      UC11([Manage Users])
      UC12([Change Password])
      UC13([Configure SMTP / App Settings])
      UC14([Manage Trucks])
      UC15([Manage Drivers])
      UC16([Manage Routes])
      UC17([Schedule / Update Collection Trip])
      UC18([Record Maintenance])
      UC19([View Fleet Dashboard])
      UC20([Send Email Notification])
    end

    %% Operator
    Operator --- UC1
    Operator --- UC2
    Operator --- UC3
    Operator --- UC4
    Operator --- UC6
    Operator --- UC9
    Operator --- UC10
    Operator --- UC12
    Operator --- UC17
    Operator --- UC19

    %% Supervisor (inherits operator)
    Supervisor --- UC1
    Supervisor --- UC2
    Supervisor --- UC3
    Supervisor --- UC4
    Supervisor --- UC5
    Supervisor --- UC6
    Supervisor --- UC7
    Supervisor --- UC8
    Supervisor --- UC9
    Supervisor --- UC10
    Supervisor --- UC12
    Supervisor --- UC15
    Supervisor --- UC16
    Supervisor --- UC17
    Supervisor --- UC18
    Supervisor --- UC19

    %% Admin (full access)
    Admin --- UC1
    Admin --- UC2
    Admin --- UC3
    Admin --- UC4
    Admin --- UC5
    Admin --- UC6
    Admin --- UC7
    Admin --- UC8
    Admin --- UC9
    Admin --- UC10
    Admin --- UC11
    Admin --- UC12
    Admin --- UC13
    Admin --- UC14
    Admin --- UC15
    Admin --- UC16
    Admin --- UC17
    Admin --- UC18
    Admin --- UC19

    %% Include / extend
    UC2 -.->|includes| UC20
    UC8 -.->|extends| UC20
    UC9 -.->|extends| UC20

    SMTP --- UC20
"""
    add_code_block(doc, mermaid_usecase, "Mermaid - flowchart (use case)")

    add_paragraph(doc, "Use-Case to Role Matrix", bold=True, size=12)
    add_table(
        doc,
        ["Use Case", "Operator", "Supervisor", "Admin"],
        [
            ["Login / Logout", "Yes", "Yes", "Yes"],
            ["Run Waste Detection", "Yes", "Yes", "Yes"],
            ["View Dashboard", "Yes", "Yes", "Yes"],
            ["View Waste History", "Yes", "Yes", "Yes"],
            ["Verify / Reject Detection", "No", "Yes", "Yes"],
            ["Export Detections", "Yes", "Yes", "Yes"],
            ["Acknowledge Alert", "No", "Yes", "Yes"],
            ["Manage Alert Rules", "No", "Yes", "Yes"],
            ["Generate PDF Report", "Yes", "Yes", "Yes"],
            ["Manage Users", "No", "No", "Yes"],
            ["Configure Settings (SMTP)", "No", "No", "Yes"],
            ["Manage Trucks", "View only", "View only", "Yes"],
            ["Manage Drivers", "No", "Edit", "Yes"],
            ["Manage Routes", "View", "Yes", "Yes"],
            ["Schedule Trip", "Own only", "Yes", "Yes"],
            ["Record Maintenance", "No", "Yes", "Yes"],
            ["Fleet Dashboard", "Yes", "Yes", "Yes"],
        ],
        col_widths=[2.6, 1.2, 1.2, 1.2],
    )
    doc.add_page_break()

    # ============================================================
    # 3. SWIM LANE DIAGRAM
    # ============================================================
    add_heading(doc, "3. Swim-lane Diagram - Waste Detection Workflow", level=1)
    add_paragraph(
        doc,
        "Swim-lanes for the end-to-end detection workflow: an operator runs a "
        "detection from the UI, the DetectionEngine performs YOLOv8 inference, "
        "the AlertManager evaluates rules, NotificationService delivers email, "
        "and a supervisor or admin verifies the record.",
    )

    mermaid_swimlane = """flowchart TB
    subgraph Operator
      O1([Open Detection Screen])
      O2([Select Image or Video])
      O3([Click Run Detection])
      O4([Review Annotated Result])
      O5([Add Notes / Save])
    end
    subgraph UI
      U1([DetectionScreen.handle_run])
      U2([Display Bin + Fill Level Boxes])
      U3([Refresh History / Dashboard])
    end
    subgraph DetectionEngine
      D1([Load YOLO Bin Model])
      D2([Load Fill-level Model])
      D3([Predict + Assign Fill Level])
      D4([Draw Annotations])
      D5([Persist Detection row])
    end
    subgraph AlertManager
      A1([Read Active AlertRules])
      A2([Count Detections in Period])
      A3{Threshold Exceeded?}
      A4([Insert Alert row])
    end
    subgraph NotificationService
      N1([Load SMTP Settings])
      N2([Send Email via STARTTLS / SSL])
    end
    subgraph Supervisor
      S1([Open Waste History])
      S2([Verify or Reject Record])
    end
    subgraph Database
      DB1[(detections)]
      DB2[(alerts)]
      DB3[(activity_logs)]
    end

    O1 --> O2 --> O3 --> U1
    U1 --> D1 --> D2 --> D3 --> D4 --> D5 --> DB1
    D5 --> U2 --> O4 --> O5
    D5 --> A1 --> A2 --> A3
    A3 -- yes --> A4 --> DB2 --> N1 --> N2
    A3 -- no --> U3
    U2 --> U3
    DB1 --> S1 --> S2 --> DB1
    S2 --> DB3
"""
    add_code_block(doc, mermaid_swimlane, "Mermaid - swim-lane (subgraph per actor)")

    add_paragraph(doc, "Lane Responsibility", bold=True, size=12)
    add_table(
        doc,
        ["Lane", "Trigger", "Outcome"],
        [
            ["Operator", "Selects image and clicks Run", "Annotated result + saved row"],
            ["UI (DetectionScreen)", "User click", "Calls DetectionEngine.detect()"],
            ["DetectionEngine", "Image path + user_id", "Detection rows + annotated JPG"],
            ["AlertManager", "After each detection", "Newly triggered alerts (deduped)"],
            ["NotificationService", "AlertRule.notify_email set", "Email delivered or error logged"],
            ["Supervisor", "Reviews pending records", "Status = verified / rejected"],
            ["Database", "All writes", "Persisted to SQLite"],
        ],
        col_widths=[1.4, 2.5, 3.0],
    )
    doc.add_page_break()

    # ============================================================
    # 4. SEQUENCE DIAGRAM (Detection + Alert)
    # ============================================================
    add_heading(doc, "4. Sequence Diagram", level=1)
    add_paragraph(
        doc,
        "Two key flows are diagrammed: (a) the login flow handled by "
        "AuthManager, and (b) the full image-detection + alert-trigger + email "
        "pipeline. Both reflect the actual call paths in the source code.",
    )

    add_paragraph(doc, "4.1 Login Sequence", bold=True, size=12)
    mermaid_seq_login = """sequenceDiagram
    actor User
    participant LS as LoginScreen
    participant AM as AuthManager
    participant DB as SQLite (users)
    participant MW as MainWindow

    User->>LS: enter username + password
    LS->>AM: login(username, password)
    AM->>DB: SELECT * FROM users WHERE username=? AND is_active=1
    DB-->>AM: row
    AM->>AM: bcrypt.checkpw(password, hash)
    alt valid
        AM->>DB: UPDATE users SET last_login=now()
        AM-->>LS: User (detached)
        LS->>MW: login_success.emit(user)
        MW->>MW: build sidebar by role
        MW->>User: show DashboardScreen
    else invalid
        AM-->>LS: None
        LS->>User: show error toast
    end
"""
    add_code_block(doc, mermaid_seq_login, "Mermaid - sequenceDiagram")

    add_paragraph(doc, "4.2 Detection + Alert Sequence", bold=True, size=12)
    mermaid_seq_detect = """sequenceDiagram
    actor Operator
    participant DS as DetectionScreen
    participant DE as DetectionEngine
    participant Y1 as YOLO Bin Model
    participant Y2 as YOLO Fill-Level Model
    participant DB as SQLite
    participant AM as AlertManager
    participant NS as NotificationService
    participant SMTP as SMTP Server

    Operator->>DS: click Run Detection (image path)
    DS->>DE: detect(image_path, user_id)
    DE->>Y1: predict(image, conf)
    Y1-->>DE: bin bounding boxes
    DE->>Y2: predict(image, conf=>=0.40)
    Y2-->>DE: fill-level boxes
    DE->>DE: IoU match + heuristic fallback
    DE->>DE: draw annotations + write result jpg
    DE->>DB: INSERT INTO detections (rows per bin)
    DB-->>DE: detection_ids
    DE-->>DS: {detections, result_image, ids}
    DS->>AM: check_alerts()
    AM->>DB: SELECT active alert_rules
    AM->>DB: COUNT detections per fill_level since period_start
    alt threshold exceeded AND not deduped
        AM->>DB: INSERT INTO alerts
        AM-->>DS: triggered list
        AM->>NS: send_email(to, subject, body)
        NS->>DB: SELECT smtp_* FROM app_settings
        NS->>SMTP: login + send_message
        SMTP-->>NS: 250 OK / error
        NS-->>AM: ok / error
    else not triggered
        AM-->>DS: []
    end
    DS->>Operator: display annotated image + toast
"""
    add_code_block(doc, mermaid_seq_detect, "Mermaid - sequenceDiagram")

    add_paragraph(doc, "4.3 Report Generation Sequence", bold=True, size=12)
    mermaid_seq_report = """sequenceDiagram
    actor User
    participant RS as ReportsScreen
    participant RE as ReportEngine
    participant AE as AnalyticsEngine
    participant DB as SQLite
    participant MPL as matplotlib
    participant RL as ReportLab

    User->>RS: choose type + date range, click Generate
    RS->>RE: generate_report(type, start, end, user_id)
    RE->>AE: get_category_distribution(start, end)
    AE->>DB: GROUP BY waste_category
    AE-->>RE: dict
    RE->>AE: get_fill_level_distribution(start, end)
    AE-->>RE: dict
    RE->>DB: SELECT detections WHERE date in range
    DB-->>RE: rows
    RE->>MPL: bar chart -> temp png
    RE->>RL: build PDF (header, tables, chart, details)
    RL-->>RE: PDF written to data/reports/
    RE->>DB: INSERT INTO reports
    RE-->>RS: file_path
    RS->>User: toast + open file
"""
    add_code_block(doc, mermaid_seq_report, "Mermaid - sequenceDiagram")
    doc.add_page_break()

    # ============================================================
    # 5. ER DIAGRAM
    # ============================================================
    add_heading(doc, "5. Entity-Relationship Diagram", level=1)
    add_paragraph(
        doc,
        "The ER diagram below maps every table created by "
        "Base.metadata.create_all(engine) in database/db_setup.py "
        "(combining models.py and fleet_models.py). PK = primary key, "
        "FK = foreign key.",
    )

    mermaid_er = """erDiagram
    USERS ||--o{ DETECTIONS : "detects (detected_by)"
    USERS ||--o{ DETECTIONS : "verifies (verified_by)"
    USERS ||--o{ ACTIVITY_LOGS : "performs"
    USERS ||--o{ ALERT_RULES : "creates"
    USERS ||--o{ ALERTS : "acknowledges"
    USERS ||--o{ REPORTS : "generates"
    USERS ||--o{ FLEET_TRIPS : "creates_by"

    ALERT_RULES ||--o{ ALERTS : "triggers"

    FLEET_TRUCKS ||--o{ FLEET_DRIVERS : "assigned"
    FLEET_TRUCKS ||--o{ FLEET_TRIPS : "runs"
    FLEET_TRUCKS ||--o{ FLEET_MAINTENANCE : "serviced"
    FLEET_DRIVERS ||--o{ FLEET_TRIPS : "drives"
    FLEET_ROUTES ||--o{ FLEET_TRIPS : "followed_by"

    USERS {
        int id PK
        string username UK
        string password_hash
        string full_name
        string email
        string role
        bool is_active
        datetime created_at
        datetime last_login
    }
    DETECTIONS {
        int id PK
        string image_path
        string result_image_path
        string waste_category
        float confidence
        string bin_fill_level
        int detected_by FK
        string status
        int verified_by FK
        datetime detected_at
        string notes
    }
    ALERT_RULES {
        int id PK
        string rule_name
        string category
        int threshold_value
        string period
        string notify_email
        bool is_active
        int created_by FK
        datetime created_at
    }
    ALERTS {
        int id PK
        int rule_id FK
        string message
        string severity
        bool is_acknowledged
        int acknowledged_by FK
        datetime triggered_at
    }
    ACTIVITY_LOGS {
        int id PK
        int user_id FK
        string action
        string details
        datetime timestamp
    }
    REPORTS {
        int id PK
        string report_type
        string file_path
        int generated_by FK
        datetime generated_at
        date date_range_start
        date date_range_end
    }
    APP_SETTINGS {
        int id PK
        string key UK
        string value
        datetime updated_at
    }
    FLEET_TRUCKS {
        int id PK
        string truck_code UK
        string plate_number UK
        float capacity
        string fuel_type
        string status
        string assigned_zone
        date purchase_date
        bool is_active
    }
    FLEET_DRIVERS {
        int id PK
        string name
        string phone
        string email
        string license_number UK
        int assigned_truck_id FK
        string status
    }
    FLEET_ROUTES {
        int id PK
        string route_name UK
        string zone
        float estimated_distance
        int estimated_duration
        string status
    }
    FLEET_TRIPS {
        int id PK
        int truck_id FK
        int driver_id FK
        int route_id FK
        datetime start_time
        datetime end_time
        float waste_weight
        string trip_status
        int created_by FK
    }
    FLEET_MAINTENANCE {
        int id PK
        int truck_id FK
        string service_type
        date service_date
        date next_service_date
        float cost
    }
"""
    add_code_block(doc, mermaid_er, "Mermaid - erDiagram")

    add_paragraph(doc, "Tables & Cardinalities", bold=True, size=12)
    add_table(
        doc,
        ["Table", "Primary Key", "Notable FKs", "Cardinality"],
        [
            ["users", "id", "-", "1 user --> * detections / logs / rules / alerts / reports / trips"],
            ["detections", "id", "detected_by, verified_by --> users",
             "* per user"],
            ["alert_rules", "id", "created_by --> users", "1 rule --> * alerts"],
            ["alerts", "id", "rule_id, acknowledged_by", "* per rule"],
            ["activity_logs", "id", "user_id --> users", "* per user"],
            ["reports", "id", "generated_by --> users", "* per user"],
            ["app_settings", "id", "(key/value store)", "Single row per key"],
            ["fleet_trucks", "id", "-", "1 truck --> * drivers / trips / maintenance"],
            ["fleet_drivers", "id", "assigned_truck_id --> trucks", "* per truck"],
            ["fleet_routes", "id", "-", "1 route --> * trips"],
            ["fleet_trips", "id", "truck_id, driver_id, route_id, created_by",
             "Join entity"],
            ["fleet_maintenance", "id", "truck_id --> trucks (CASCADE)",
             "* per truck"],
        ],
        col_widths=[1.4, 0.8, 2.4, 2.7],
    )
    doc.add_page_break()

    # ============================================================
    # 6. DFD - LEVEL 0 (Context) + LEVEL 1
    # ============================================================
    add_heading(doc, "6. Data Flow Diagrams (DFD)", level=1)

    # Level 0
    add_paragraph(doc, "6.1 DFD Level 0 (Context Diagram)", bold=True, size=12)
    add_paragraph(
        doc,
        "Single process (SWMS) with external entities: Operator, Supervisor, "
        "Administrator, SMTP Server, and the local file system (image / "
        "report storage).",
    )

    mermaid_dfd_l0 = """flowchart LR
    Operator((Operator))
    Supervisor((Supervisor))
    Admin((Administrator))
    SMTP[[SMTP Server]]
    FS[[Local File System]]

    Operator -- "image upload, credentials, run-detect" --> SWMS((Smart Waste Management System))
    SWMS -- "annotated image, history, alerts" --> Operator

    Supervisor -- "verify, acknowledge, manage rules" --> SWMS
    SWMS -- "pending lists, KPIs, reports" --> Supervisor

    Admin -- "user CRUD, SMTP config, fleet CRUD" --> SWMS
    SWMS -- "user list, audit logs, fleet KPIs" --> Admin

    SWMS -- "MIME email" --> SMTP
    SMTP -- "delivery status" --> SWMS

    SWMS -- "save annotated jpg / PDF report" --> FS
    FS -- "read image, render report" --> SWMS
"""
    add_code_block(doc, mermaid_dfd_l0, "Mermaid - DFD Level 0")

    # Level 1
    add_paragraph(doc, "6.2 DFD Level 1 (Process Decomposition)", bold=True, size=12)
    add_paragraph(
        doc,
        "Decomposes SWMS into its main processes and the data stores (SQLite "
        "tables) they read/write.",
    )

    mermaid_dfd_l1 = """flowchart TB
    %% External entities
    Op((Operator))
    Sup((Supervisor))
    Adm((Administrator))
    SMTP[[SMTP Server]]
    FS[[File System])]

    %% Processes
    P1([1. Authenticate User])
    P2([2. Detect Waste])
    P3([3. Evaluate Alerts])
    P4([4. Generate Analytics / Charts])
    P5([5. Generate PDF Report])
    P6([6. Manage Users])
    P7([7. Manage Alert Rules])
    P8([8. Manage Fleet & Trips])
    P9([9. Send Notification])
    P10([10. Log Activity])

    %% Data stores
    DS1[(D1: users)]
    DS2[(D2: detections)]
    DS3[(D3: alert_rules)]
    DS4[(D4: alerts)]
    DS5[(D5: activity_logs)]
    DS6[(D6: reports)]
    DS7[(D7: app_settings)]
    DS8[(D8: fleet_trucks / drivers / routes / trips / maintenance)]

    %% Authentication
    Op -- credentials --> P1
    Sup -- credentials --> P1
    Adm -- credentials --> P1
    P1 <-- "hash / role lookup" --> DS1
    P1 -- session user --> Op
    P1 -- audit event --> P10

    %% Detection
    Op -- image path --> P2
    FS -- image bytes --> P2
    P2 -- annotated jpg --> FS
    P2 -- detection row --> DS2
    P2 -- "new detection" --> P3

    %% Alert evaluation
    P3 <-- read rules --> DS3
    P3 <-- count detections --> DS2
    P3 -- alert row --> DS4
    P3 -- email content --> P9

    %% Notification
    P9 <-- "SMTP credentials" --> DS7
    P9 -- MIME --> SMTP

    %% Analytics + reports
    Op -- view dashboard --> P4
    Sup -- view dashboard --> P4
    P4 <-- aggregates --> DS2
    P4 <-- aggregates --> DS4
    P4 -- KPIs / charts --> Op
    Sup -- "type + date range" --> P5
    P5 <-- read --> DS2
    P5 -- pdf --> FS
    P5 -- report row --> DS6

    %% User mgmt
    Adm -- user CRUD --> P6
    P6 <-- read/write --> DS1
    P6 -- audit --> P10

    %% Rule mgmt
    Sup -- rule CRUD --> P7
    Adm -- rule CRUD --> P7
    P7 <-- read/write --> DS3
    P7 -- audit --> P10

    %% Fleet
    Adm -- truck/driver/route/trip CRUD --> P8
    Sup -- driver/route/trip/maintenance --> P8
    Op -- "own trip status" --> P8
    P8 <-- read/write --> DS8
    P8 -- audit --> P10

    %% Logging
    P10 -- append --> DS5
"""
    add_code_block(doc, mermaid_dfd_l1, "Mermaid - DFD Level 1")

    add_paragraph(doc, "Data Store Summary", bold=True, size=12)
    add_table(
        doc,
        ["Data Store", "Table(s)", "Written by", "Read by"],
        [
            ["D1 users", "users", "P1, P6", "P1, P4, P5, P6, P10"],
            ["D2 detections", "detections", "P2, supervisor verify",
             "P2, P3, P4, P5"],
            ["D3 alert_rules", "alert_rules", "P7", "P3, P7"],
            ["D4 alerts", "alerts", "P3", "P3, P4"],
            ["D5 activity_logs", "activity_logs", "P10", "Admin via Logs UI"],
            ["D6 reports", "reports", "P5", "P5 (list)"],
            ["D7 app_settings", "app_settings", "P9 (via Admin), init_db",
             "P2 (threshold), P9 (SMTP)"],
            ["D8 fleet", "fleet_trucks, fleet_drivers, fleet_routes, "
             "fleet_trips, fleet_maintenance", "P8",
             "P4 (fleet dashboard), P8"],
        ],
        col_widths=[1.4, 2.6, 1.4, 1.6],
    )

    # ---- Appendix ----
    doc.add_page_break()
    add_heading(doc, "Appendix - How to render the diagrams", level=1)
    add_paragraph(
        doc,
        "1. Open https://mermaid.live (or use the Mermaid VS Code extension).\n"
        "2. Paste any of the code blocks above into the editor.\n"
        "3. Export as PNG / SVG and embed back into this document if needed.\n\n"
        "All diagram source code is plain text Mermaid, so the diagrams remain "
        "version-controllable alongside the source repository.",
    )

    doc.save(out_path)
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    out = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "SWMS_System_Diagrams.docx",
    )
    build_document(out)
